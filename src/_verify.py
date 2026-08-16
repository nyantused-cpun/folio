# -*- coding: utf-8 -*-
import os
import re
import math

SCRIPT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # 仓库根

MIN_PPTX_SIZE = 5 * 1024
MIN_HTML_SIZE = 500
MIN_DOCX_SIZE = 5 * 1024
MIN_XLSX_SIZE = 1 * 1024
MIN_YAML_SIZE = 50

# L4: 去 AI 痕迹（单一数据源：_paths.py）
from _paths import BANNED_WORDS, BANNED_PHRASES


def verify_pptx(path):
    if not os.path.exists(path):
        return False, f"文件不存在: {path}"
    if not path.lower().endswith(".pptx"):
        return False, f"不是 .pptx 文件: {path}"

    size = os.path.getsize(path)
    if size < MIN_PPTX_SIZE:
        return False, f"文件过小 ({size} bytes, 最少 {MIN_PPTX_SIZE} bytes): {path}"

    try:
        from pptx import Presentation
        prs = Presentation(path)
    except Exception as e:
        return False, f"无法打开 PPTX: {e}"

    slide_count = len(prs.slides)
    if slide_count < 1:
        return False, "PPTX 没有幻灯片"

    return True, f"OK ({slide_count} 页, {size} bytes)"


def verify_html(path):
    if not os.path.exists(path):
        return False, f"文件不存在: {path}"
    if not (path.lower().endswith(".html") or path.lower().endswith(".htm")):
        return False, f"不是 HTML 文件: {path}"

    size = os.path.getsize(path)
    if size < MIN_HTML_SIZE:
        return False, f"文件过小 ({size} bytes, 最少 {MIN_HTML_SIZE} bytes): {path}"

    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return False, f"无法读取 HTML: {e}"

    if "<!DOCTYPE html>" not in content and "<!doctype html>" not in content:
        return False, "缺少 <!DOCTYPE html> 声明"

    if "</html>" not in content:
        return False, "缺少 </html> 闭合标签"

    # 放宽检查：只要有 CSS 样式引用即可（Tailwind/内联/外部 CSS 均可）
    has_css = (
        "<style" in content or
        "tailwindcss" in content or
        'rel="stylesheet"' in content or
        ".css" in content
    )
    if not has_css:
        return False, "未检测到任何 CSS 样式引用"

    msg = f"OK ({size} bytes)"
    # 美学观察（D-120，观察模式不阻断）：渲染后几何美学检查
    # （重叠/字号层级/溢出/大空白/对比度）；playwright 不可用时静默跳过
    try:
        from _density import check_density_html
        findings = check_density_html(path) or []
        real = [f for f in findings
                if not f.startswith(("playwright", "密度检查执行失败"))]
        if real:
            msg += f" | 美学观察 {len(real)} 条: " + "; ".join(real[:3])
    except Exception:
        pass
    return True, msg


def verify_docx(path, client_name=""):
    if not os.path.exists(path):
        return False, f"文件不存在: {path}"
    if not path.lower().endswith(".docx"):
        return False, f"不是 .docx 文件: {path}"

    size = os.path.getsize(path)
    if size < MIN_DOCX_SIZE:
        return False, f"文件过小 ({size} bytes, 最少 {MIN_DOCX_SIZE} bytes): {path}"

    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        return False, f"无法打开 DOCX: {e}"

    para_count = len(doc.paragraphs)
    if para_count < 1:
        return False, "DOCX 没有段落内容"

    non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
    if non_empty == 0:
        return False, "DOCX 所有段落均为空"

    # L4: 去 AI 痕迹（不阻断，仅警告）
    l4_ok, l4_issues = check_ai_traces(path)
    l4_msg = ""
    if not l4_ok:
        l4_msg = " | L4警告: " + "; ".join(l4_issues[:3])

    # L5: 废标风险（不阻断，仅警告）
    l5_ok, l5_issues = check_bid_risks(path, client_name)
    l5_msg = ""
    if not l5_ok:
        l5_msg = " | L5风险: " + "; ".join(l5_issues[:3])

    return True, f"OK ({para_count} 段落, {non_empty} 非空, {size} bytes){l4_msg}{l5_msg}"


# ============================================================
# L4: 去 AI 痕迹检查
# ============================================================

def check_ai_traces(path):
    """检查 docx 中的 AI 痕迹。

    三项检查：
    1. burstiness（句子长度方差，AI 文本方差低）
    2. 禁用词（赋能/抓手/闭环...）
    3. 模板句式（综上所述/总而言之...）

    返回 (pass: bool, issues: list[str])
    """
    from docx import Document
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if not full_text:
        return True, []

    issues = []

    # 1. burstiness 检查
    sentences = re.split(r"[。！？.!?]", full_text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 5]
    if len(sentences) >= 10:
        lengths = [len(s) for s in sentences]
        mean_len = sum(lengths) / len(lengths)
        variance = sum((length - mean_len) ** 2 for length in lengths) / len(lengths)
        std_dev = math.sqrt(variance)
        burstiness = std_dev / mean_len if mean_len > 0 else 0
        # burstiness < 0.3 说明句子长度太均匀，疑似 AI 生成
        if burstiness < 0.3:
            issues.append(f"burstiness={burstiness:.2f}（偏低，句子长度过于均匀，疑似AI生成）")

    # 2. 禁用词检查
    for word in BANNED_WORDS:
        count = full_text.count(word)
        if count > 0:
            issues.append(f"禁用词「{word}」出现 {count} 次")

    # 3. 模板句式检查
    for phrase in BANNED_PHRASES:
        count = full_text.count(phrase)
        if count > 0:
            issues.append(f"模板句式「{phrase}」出现 {count} 次")

    return len(issues) == 0, issues


# ============================================================
# L5: 废标风险检查
# ============================================================

def check_bid_risks(path, client_name=""):
    """检查标书废标风险。

    三项检查：
    1. 评分覆盖率：每项评分是否都有对应章节
    2. 资质提及：must_mention 中的资质是否出现
    3. 格式合规：检查页边距/字体/暗标等

    返回 (pass: bool, issues: list[str])
    """
    import json

    issues = []

    # 加载 bid_criteria.json
    criteria_path = None
    if client_name:
        criteria_path = os.path.join(SCRIPT_DIR, "output", client_name, "bid_criteria.json")
    if not criteria_path or not os.path.exists(criteria_path):
        return True, []  # 无招标标准，跳过

    with open(criteria_path, "r", encoding="utf-8") as f:
        try:
            criteria = json.load(f)
        except json.JSONDecodeError as e:
            return False, [f"bid_criteria.json 解析失败: {e}"]

    from docx import Document
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # 1. 评分覆盖率
    scoring = criteria.get("scoring", [])
    uncovered = []
    for item in scoring:
        item_name = item.get("item", "")
        if item_name and item_name not in full_text:
            # 检查子项
            sub_items = item.get("sub_items", [])
            if sub_items:
                sub_covered = any(
                    sub.get("name", "") in full_text for sub in sub_items
                )
                if not sub_covered:
                    uncovered.append(item_name)
            else:
                uncovered.append(item_name)
    if uncovered:
        issues.append(f"评分项未覆盖: {', '.join(uncovered)}")

    # 2. 资质提及
    quals = criteria.get("qualifications", {})
    must_mention = quals.get("must_mention", [])
    missing_quals = [q for q in must_mention if q not in full_text]
    if missing_quals:
        issues.append(f"资质未提及: {', '.join(missing_quals)}")

    # 3. 格式合规（简单检查）
    fmt = criteria.get("format_requirements", {})
    if fmt.get("dark_bid"):
        # 暗标：检查是否含公司名/作者信息
        # 这里简单检查，实际应更严格
        author_fields = ["作者", "撰写人", "编制人", "公司名"]
        for field in author_fields:
            if field in full_text:
                issues.append(f"暗标模式但含「{field}」信息")
                break

    return len(issues) == 0, issues


def verify_xlsx(path):
    if not os.path.exists(path):
        return False, f"文件不存在: {path}"
    if not path.lower().endswith(".xlsx"):
        return False, f"不是 .xlsx 文件: {path}"

    size = os.path.getsize(path)
    if size < MIN_XLSX_SIZE:
        return False, f"文件过小 ({size} bytes, 最少 {MIN_XLSX_SIZE} bytes): {path}"

    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
    except Exception as e:
        return False, f"无法打开 XLSX: {e}"

    sheet_count = len(wb.sheetnames)
    if sheet_count < 1:
        return False, "XLSX 没有工作表"

    return True, f"OK ({sheet_count} 工作表, {size} bytes)"


def verify_yaml(path):
    if not os.path.exists(path):
        return False, f"文件不存在: {path}"
    if not (path.lower().endswith(".yml") or path.lower().endswith(".yaml")):
        return False, f"不是 YAML 文件: {path}"

    size = os.path.getsize(path)
    if size < MIN_YAML_SIZE:
        return False, f"文件过小 ({size} bytes, 最少 {MIN_YAML_SIZE} bytes): {path}"

    try:
        import yaml
        with open(path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f)
    except Exception as e:
        return False, f"无法解析 YAML: {e}"

    if data is None:
        return False, "YAML 文件内容为空或仅含注释"

    # v2.0 版式维（§10）：spec 文件（含 pages）追加 6 条版式检查
    if isinstance(data, dict) and isinstance(data.get("pages"), list):
        lay_ok, lay_errors, lay_warnings = verify_spec_layout_data(data)
        msg = f"OK ({type(data).__name__}, {size} bytes)"
        if lay_warnings:
            msg += f" | 版式警告 {len(lay_warnings)} 条: " + "; ".join(lay_warnings[:3])
        if lay_errors:
            return False, (f"{msg} | 版式校验 {len(lay_errors)} 条错误: "
                           + "; ".join(lay_errors[:3]))
        return True, msg

    return True, f"OK ({type(data).__name__}, {size} bytes)"


# ============================================================
# v2.0 版式维（dev_plan_visual_v2_2026-07-25 §10，6 条检查）
# ============================================================

# 结论式标题量化词（§10 检查 3：含数字或量化词）
_QUANT_WORDS = ("%", "翻倍", "倍", "降至", "提升", "提高", "降低", "减少",
                "增长", "覆盖", "节约", "缩短", "万", "亿")


def check_type_role_consistency(spec):
    """A-2 散值字号统计（观察模式）：扫描全 spec 元素的 fontSize 散值。

    同一 spec 内同一散值字号出现 >2 次 → 提示收敛为 type_roles 角色。
    只报告不阻断（观察模式，exit code 不变）。
    """
    from collections import Counter
    counts = Counter()
    for page in spec.get("pages", []) or []:
        if not isinstance(page, dict):
            continue
        for elem in page.get("elements", []) or []:
            if not isinstance(elem, dict):
                continue
            fs = elem.get("fontSize")
            if isinstance(fs, (int, float)) and not isinstance(fs, bool):
                counts[fs] += 1
    return [f"[观察] 散值字号 {fs} 出现 {n} 次，建议收敛为 type_roles 角色"
            for fs, n in sorted(counts.items()) if n > 2]


def check_shell_repetition(spec):
    """A-3 反同壳检查（观察模式）：相邻 ≥3 页 layout + 构件指纹相同 → [观察]。

    指纹 = layout + 构件类型序列 + diagram subtype 序列（文案内容不参与）。
    只报告不阻断（观察模式，exit code 不变）。
    """
    pages = [p for p in spec.get("pages", []) or [] if isinstance(p, dict)]
    fingerprints = []
    for p in pages:
        layout = p.get("layout", "")
        types, subtypes = [], []
        for e in p.get("elements", []) or []:
            if not isinstance(e, dict):
                continue
            types.append(e.get("type", ""))
            if e.get("type") == "diagram":
                subtypes.append(f"{e.get('diagram_type', '')}/{e.get('subtype', '')}")
        fingerprints.append((layout, tuple(types), tuple(subtypes)))

    msgs = []
    i, n = 0, len(fingerprints)
    while i < n:
        j = i
        while j < n and fingerprints[j] == fingerprints[i]:
            j += 1
        if j - i >= 3:
            layout, types, subtypes = fingerprints[i]
            parts = [layout] if layout else ["无版式"]
            if types:
                parts.append("构件:" + "+".join(types))
            if subtypes:
                parts.append("图:" + "+".join(subtypes))
            msgs.append(f"[观察] 页面 {i + 1}-{j} 构图壳重复（{'；'.join(parts)}）")
        i = j
    return msgs


def check_evidence_ledger(spec):
    """B-3 证据台账编号唯一性（观察模式）：num 非空且重复 → [观察]。

    只报告不阻断（观察模式，exit code 不变）。
    """
    from collections import Counter
    msgs = []
    for pi, page in enumerate(spec.get("pages", []) or []):
        if not isinstance(page, dict):
            continue
        for ei, elem in enumerate(page.get("elements", []) or []):
            if not isinstance(elem, dict) or elem.get("type") != "evidence_ledger":
                continue
            nums = [it.get("num", "") for it in elem.get("items", []) or []
                    if isinstance(it, dict) and it.get("num")]
            dup = [n for n, c in Counter(nums).items() if c > 1]
            if dup:
                msgs.append(f"[观察] pages[{pi}].elements[{ei}]: "
                            f"evidence_ledger 证据编号重复 {', '.join(dup)}")
    return msgs


# D-122：构图母板 → 期望构件特征（讲法观察：声明了母板但页面构件不匹配）
_COMPOSITION_HINTS = {
    "full_claim": ("hero/action_title 类大标题构件", {"hero", "action_title"}),
    "editorial_columns": ("图文穿插类（info_cards/pullquote）",
                          {"info_cards", "pullquote", "view_cards"}),
    "architecture_board": ("diagram architecture 类图", None),
    "evidence_ledger": ("证据结构（evidence_ledger/info_cards）",
                        {"evidence_ledger", "info_cards"}),
    "flow_spine": ("diagram flow 类流程图", None),
    "scenario_sequence": ("多段 section_tag/action_title 序列",
                          {"section_tag", "action_title"}),
    "data_narrative": ("数据元素（stat_cards/kpi_cards/table）",
                       {"stat_cards", "kpi_cards", "table"}),
    "product_simulation": ("仿真元素（simulation 组件）", {"simulation"}),
    "timeline_gantt": ("diagram timeline 类（module_gantt/milestone_gantt）", None),
    "comparison_matrix": ("对比类（table/fit_gap/cbm/raci/crud/quadrant）",
                          {"table", "diagram", "qa_block"}),
    "decision_board": ("决策类（callout_block/决策图）", {"callout_block", "diagram"}),
    "capability_graph": ("能力图（capability_map/platform_hub）", None),
}

# diagram 类母板的 subtype 命中表
_COMPOSITION_DIAGRAM_SUBTYPES = {
    "architecture_board": ("4a", "layered", "integration", "biz_overview",
                           "deployment", "platform_hub", "pyramid"),
    "flow_spine": ("sequence", "swimlane", "cross_system", "parallel",
                   "decision", "flow_rows"),
    "timeline_gantt": ("horizontal", "vertical", "module_gantt", "milestone_gantt"),
    "capability_graph": ("capability_map", "platform_hub"),
}


def check_composition_fit(spec):
    """D-122 讲法观察：页面声明 composition 母板后，构件是否匹配母板预期。

    只报告不阻断（观察模式）：declared 母板但页面构件不匹配 → [观察] 提示。
    目的：把「该对比矩阵的用了卡片墙」这类讲法错误从靠 AI 自觉变成机械提示。
    """
    msgs = []
    for pi, page in enumerate(spec.get("pages", []) or []):
        if not isinstance(page, dict):
            continue
        comp = page.get("composition")
        if not comp:
            continue
        if isinstance(comp, str):
            comp = [comp]
        elem_types = set()
        subtypes = set()
        for e in page.get("elements", []) or []:
            if not isinstance(e, dict):
                continue
            elem_types.add(e.get("type", ""))
            if e.get("type") == "diagram":
                subtypes.add(e.get("subtype", ""))
        for c in comp:
            hint, type_set = _COMPOSITION_HINTS.get(c, ("", None))
            if hint == "":
                continue  # 非法枚举由 schema error 报，观察项不重复
            if type_set is None:
                want = set(_COMPOSITION_DIAGRAM_SUBTYPES.get(c, ()))
                hit = bool(subtypes & want)
                if not hit:
                    msgs.append(
                        f"[观察] pages[{pi}] 声明 {c} 母板但无对应图 subtype "
                        f"（期望 {'/'.join(want)}，实际 {'/'.join(sorted(subtypes)) or '无'}）")
            else:
                hit = bool(elem_types & type_set)
                if not hit:
                    msgs.append(
                        f"[观察] pages[{pi}] 声明 {c} 母板但缺 {hint} "
                        f"（现有构件：{'+'.join(sorted(elem_types)) or '无'}）")
    return msgs


def verify_spec_layout_data(spec):
    """spec 版式维 6 条检查（§10）。返回 (ok, errors, warnings)。

    1 版式合法（error）：layout ∈ P01-P11 ∪ 旧值白名单 + 必需构件齐全
      ——schema.validate_layout_errors 复核（单点在 schema，此处不重复实现）
    2 主题合法（error）：theme ∈ THEMES + v2 spec 无 hex 字面量
      ——schema._validate_doc_theme / _validate_spec_hex 复核
    3 结论式标题（warning）：action_title 全文 ≥12 字且含数字或量化词
    4 图例完备（error）：capability_map（三态语义色 >2）同页无 legend_bar；
      flow_rows roles>2 走 §8.1 自动 legend，豁免
    5 虚线纪律（warning）：dashed_opt 行内 badge 等——schema 容量/纪律
      warnings 复核（单点在 schema.validate_element_warnings）
    6 容量阈值（warning）：flow_rows 行/卡、P04 pain、P07 table 行数等
      ——schema.validate_element_warnings / validate_layout_warnings 复核
    """
    from _renderer import schema as _schema
    errors, warnings = [], []

    # 检查 1 + 2（schema 复核：版式/theme/hex 的 error 全集）
    errors.extend(_schema.validate_spec(spec))

    for pi, page in enumerate(spec.get("pages", []) or []):
        if not isinstance(page, dict):
            continue
        elements = [e for e in (page.get("elements", []) or [])
                    if isinstance(e, dict)]
        where = f"pages[{pi}]"
        # 检查 6（版式容量）+ 顺序
        warnings.extend(_schema.validate_layout_warnings(page, page_index=pi))
        # 检查 3：结论式标题
        for ei, elem in enumerate(elements):
            if elem.get("type") != "action_title":
                continue
            _, normalized = _schema._proto.normalize_element(elem)
            full = "".join(seg.get("t", "") for seg in normalized["segments"])
            has_num = any(ch.isdigit() for ch in full)
            has_quant = any(w in full for w in _QUANT_WORDS)
            if len(full) < 12 or not (has_num or has_quant):
                warnings.append(
                    f"{where}.elements[{ei}]: action_title 不满足结论式标题"
                    f"（需 ≥12 字且含数字或量化词），实际 {len(full)} 字")
        # 检查 4：图例完备
        has_legend = any(e.get("type") == "legend_bar" for e in elements)
        for ei, elem in enumerate(elements):
            if elem.get("type") != "diagram":
                continue
            # 检查 5 + 6（元素级容量/纪律）
            warnings.extend(_schema.validate_element_warnings(elem, index=ei))
            if elem.get("subtype") == "capability_map" and not has_legend:
                errors.append(
                    f"{where}.elements[{ei}]: capability_map 三态语义色 >2 "
                    "但同页无 legend_bar（图例完备，§10）")

        # 检查 9a：AI 配色纪律（v2 新增，Kimi 方法论：单页语义色 ≤4；
        # 禁红/紫/黄/绿高饱和四色同页——AI 经典配色）
        semantic = set()
        for elem in elements:
            etype = elem.get("type")
            if etype == "stat_cards":
                for c in elem.get("cards", []) or []:
                    if isinstance(c, dict) and c.get("tone") in ("lit", "part", "gap"):
                        semantic.add(c["tone"])
            elif etype == "legend_bar":
                for it in elem.get("items", []) or []:
                    if isinstance(it, dict) and it.get("swatch"):
                        semantic.add(str(it["swatch"]))
            elif etype == "pain_cards":
                for c in elem.get("cards", []) or []:
                    if isinstance(c, dict) and c.get("level") in ("P0", "P1", "P2"):
                        semantic.add(c["level"])
            elif etype == "action_title":
                for seg in elem.get("segments", []) or []:
                    if isinstance(seg, dict) and seg.get("hl") in ("yellow", "red", "green"):
                        semantic.add(f'hl_{seg["hl"]}')
            elif etype == "diagram" and elem.get("subtype") == "flow_rows":
                roles = elem.get("roles") or {}
                if isinstance(roles, dict):
                    semantic.update(f"role_{k}" for k in roles)
        if len(semantic) > 4:
            warnings.append(f"{where}: 单页语义色 {len(semantic)} 种 >4"
                            "（AI 配色纪律，§10 检查 9）")
        families = set()
        for s in semantic:
            if s in ("role_legal", "P0", "hl_red"):
                families.add("red")
            elif s == "role_sys":
                families.add("purple")
            elif s == "hl_yellow":
                families.add("yellow")
            elif s in ("role_fin", "lit", "hl_green"):
                families.add("green")
        if families == {"red", "purple", "yellow", "green"}:
            warnings.append(f"{where}: 红/紫/黄/绿高饱和四色同页"
                            "（Kimi 禁 AI 经典配色，§10 检查 9）")

    # A-2 散值字号观察 + A-3 反同壳观察 + B-3 证据编号重复观察
    # + D-122 讲法观察（观察模式）
    warnings.extend(check_type_role_consistency(spec))
    warnings.extend(check_shell_repetition(spec))
    warnings.extend(check_evidence_ledger(spec))
    warnings.extend(check_composition_fit(spec))

    return len(errors) == 0, errors, warnings


def verify_all(paths_dict):
    results = {}
    all_ok = True

    verifiers = {
        "pptx": (verify_pptx, ".pptx"),
        "html": (verify_html, ".html"),
        "docx": (verify_docx, ".docx"),
        "xlsx": (verify_xlsx, ".xlsx"),
        "yaml": (verify_yaml, ".yml"),
    }

    for key, (fn, _ext) in verifiers.items():
        path = paths_dict.get(key)
        if not path:
            continue
        ok, msg = fn(path)
        results[key] = {"ok": ok, "msg": msg, "path": path}
        if not ok:
            all_ok = False

    return all_ok, results


def auto_verify(path, client_name=""):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pptx":
        return verify_pptx(path)
    elif ext in (".html", ".htm"):
        return verify_html(path)
    elif ext == ".docx":
        return verify_docx(path, client_name)
    elif ext == ".xlsx":
        return verify_xlsx(path)
    elif ext == ".pptd":
        # pptd 主文件（yaml 内容）：格式检查 + 版式维复检（§10 检查 7/8/9b）
        if not os.path.exists(path):
            return False, f"文件不存在: {path}"
        try:
            import yaml
            with open(path, "r", encoding="utf-8") as f:
                deck = yaml.safe_load(f)
        except Exception as e:
            return False, f"无法解析 pptd（yaml）: {e}"
        if not isinstance(deck, dict) or "pages" not in deck:
            return False, "pptd 缺 pages 结构"
        ok, msg = True, f"OK (pptd, {os.path.getsize(path)} bytes)"
        lay_ok, lay_errors, lay_warnings = verify_pptd_deck(
            os.path.dirname(os.path.abspath(path)) or ".")
        if lay_warnings:
            msg += f" | 版式警告 {len(lay_warnings)} 条: " + "; ".join(lay_warnings[:3])
        if lay_errors:
            return False, (f"{msg} | pptd 版式维 {len(lay_errors)} 条错误: "
                           + "; ".join(lay_errors[:3]))
        return ok, msg
    elif ext in (".yml", ".yaml"):
        return verify_yaml(path)
    else:
        return False, f"不支持的文件类型: {ext}"


# ============================================================
# v2.0 pptd 版式维复检（§10 检查 7/8/9b，v2 修订新增）
# ============================================================

# 检查 7 五类 Warning 类型：单点定义在 _pptd.py（checker 封装层）
# （局部 import 见 verify_pptd_deck，避免模块级依赖）

# 检查 9b 卡类 elementId（去唯一化前缀 v2{y}-/dg{y}- 后匹配）
_CARD_ID_RE = re.compile(
    r"^(?:(?:v2\d+|dg\d+)-)?(?:(?:stc|kpi|pain|info)-\d+|fr-c\d+-\d+|hero-sc\d+)$")


def _bounds_intersect(a, b, shrink=0.0):
    """矩形相交判定；shrink 为 a 的四边内缩容差（穿字检查防边缘误报）。"""
    ax1, ay1 = a[0] + shrink, a[1] + shrink
    ax2, ay2 = a[0] + a[2] - shrink, a[1] + a[3] - shrink
    return not (ax2 <= b[0] or b[0] + b[2] <= ax1
                or ay2 <= b[1] or b[1] + b[3] <= ay1)


def _bounds_contains(outer, inner):
    return (outer[0] <= inner[0] and outer[1] <= inner[1]
            and outer[0] + outer[2] >= inner[0] + inner[2]
            and outer[1] + outer[3] >= inner[1] + inner[3])


def verify_pptd_deck(deck_dir):
    """pptd 工程版式维复检（§10 检查 7/8/9b）。返回 (ok, errors, warnings)。

    - 检查 7（error）：check_report.json（pptd-build 时 _pptd.py 写入）中
      五类 Warning 计数 >0；无报告文件时跳过（老工程兼容）
    - 检查 8（warning）：连接线（*Connector）bounds 与任何 text 元素
      bounds 相交（Kimi 制图纪律：连接线止于节点边缘、绝不穿越文字）
    - 检查 9b（warning）：卡片嵌套 >1 层（卡类 shape 完全包含另一卡类）
    """
    import json
    import yaml
    from _pptd import PPTD_WARN_TYPES  # 单点定义（checker 封装层）
    errors, warnings = [], []

    # ---- 检查 7：checker 五类 warning 复检 ----
    report_path = os.path.join(deck_dir, "check_report.json")
    if os.path.exists(report_path):
        try:
            with open(report_path, "r", encoding="utf-8") as f:
                check_report = json.load(f)
        except (OSError, ValueError) as e:
            warnings.append(f"check_report.json 解析失败: {e}")
            check_report = {}
        warn_counts = check_report.get("warnings", {})
        bad = {t: n for t, n in warn_counts.items()
               if t in PPTD_WARN_TYPES and n > 0}
        if bad:
            detail = ", ".join(f"{t}×{n}" for t, n in sorted(bad.items()))
            errors.append(f"pptd checker 五类 Warning 未处理：{detail}（§10 检查 7）")

    # ---- 检查 8/9b：pages/*.page 几何扫描 ----
    pages_dir = os.path.join(deck_dir, "pages")
    page_files = []
    if os.path.isdir(pages_dir):
        page_files = [os.path.join(pages_dir, fn) for fn in os.listdir(pages_dir)
                      if fn.endswith(".page")]
    elif os.path.exists(report_path) or any(
            fn.endswith(".pptd") for fn in (os.listdir(deck_dir)
                                            if os.path.isdir(deck_dir) else [])):
        pass  # 有工程但无 pages 目录（空 deck），不报错
    for pf in sorted(page_files):
        try:
            with open(pf, "r", encoding="utf-8") as f:
                page = yaml.safe_load(f)
        except (OSError, ValueError):
            continue
        elements = (page or {}).get("elements", []) or []
        pname = os.path.basename(pf)
        # D-2 图形白名单（观察模式）：shape preset / connector kind / diagram 无图像化
        warnings.extend(_diagram_shape_whitelist_warnings(elements, pname))
        connectors = [e for e in elements
                      if e.get("elementType") == "shape"
                      and "Connector" in str(e.get("shapeName", ""))]
        texts = [e for e in elements if e.get("elementType") == "text"]
        for conn in connectors:
            cb = conn.get("bounds")
            if not cb or len(cb) != 4:
                continue
            for t in texts:
                tb = t.get("bounds")
                if not tb or len(tb) != 4:
                    continue
                if _bounds_intersect(cb, tb, shrink=2.0):
                    warnings.append(
                        f"{pname}: 连接线 {conn.get('elementId')} 穿过文本 "
                        f"{t.get('elementId')}（连接线止于节点边缘，§10 检查 8）")
        # 卡片嵌套：卡类完全包含另一卡类（装饰件/容器底不在卡类正则内）
        cards = [e for e in elements
                 if e.get("elementType") == "shape"
                 and _CARD_ID_RE.match(str(e.get("elementId", "")))
                 and e.get("bounds") and len(e["bounds"]) == 4]
        for outer in cards:
            for inner in cards:
                if outer is inner:
                    continue
                if _bounds_contains(outer["bounds"], inner["bounds"]):
                    warnings.append(
                        f"{pname}: 卡片嵌套 {outer.get('elementId')} ⊃ "
                        f"{inner.get('elementId')}（同页卡片不嵌套，§10 检查 9）")

    return len(errors) == 0, errors, warnings


def _diagram_shape_whitelist_warnings(elements, pname):
    """D-2 图形白名单（观察模式）：shape 名在注册 preset/connector 表内、
    diagram 区域无图像化。只报告不阻断（§2/§8 观察模式纪律）。

    白名单单点来源 = _pptd_convert 的 _SHAPE_NAME_MAP / _CONNECTOR_MAP
    （转换器支持什么就允许什么，方言层与转换层不分叉）。
    """
    from _pptd_convert import _SHAPE_NAME_MAP, _CONNECTOR_MAP
    allowed = set(_SHAPE_NAME_MAP) | set(_CONNECTOR_MAP)
    msgs = []
    for e in elements:
        et = e.get("elementType")
        if et == "shape":
            sn = str(e.get("shapeName", ""))
            if not sn:
                continue
            if "Connector" in sn:
                if sn not in _CONNECTOR_MAP:
                    msgs.append(f"{pname}: 连接线 kind 未注册 {e.get('elementId')} "
                                f"shapeName={sn}（合法: {sorted(_CONNECTOR_MAP)}）")
            elif sn not in allowed:
                msgs.append(f"{pname}: shape 预设名不在白名单 {e.get('elementId')} "
                            f"shapeName={sn}")
        elif et == "image":
            eid = str(e.get("elementId", ""))
            if eid.startswith("dg"):
                msgs.append(f"{pname}: diagram 图形区域出现图片元素 {eid}"
                            "（图像化违规，图应走原生 shape/connector）")
    return [f"[观察] {m}" for m in msgs]



