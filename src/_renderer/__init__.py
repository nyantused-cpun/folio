# -*- coding: utf-8 -*-
"""统一渲染器：spec 驱动的 HTML/PPT/DOCX 生成。"""

import os
import json
import yaml
from string import Template

from _renderer.elements import (
    RenderReport, _esc, _nl2br, degrade_text, empty_payload_reason,
    is_empty_payload, normalize_element)
from _renderer.theme import resolve_theme


class NotInvokedViaCLIError(Exception):
    """直接 import _renderer 被阻断时抛出。"""
    pass


class RenderBlockedError(Exception):
    """spec.confirmed 未设为 true 时抛出。"""
    pass


class OutputPathNotAllowedError(Exception):
    """输出路径不在白名单内时抛出。"""
    pass


SCRIPT_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))  # 仓库根（output 白名单）

# outline-to-spec 提取失败占位页的标记串（生成点：_outline_to_spec.py
# 失败章节的占位 text 元素 content），Renderer init 据此进 report.warnings
_EXTRACT_FAIL_MARKER = "（本章节内容提取失败"


def _validate_output_path(output_path):
    """验证输出路径是否在允许的白名单内。

    用 commonpath 做目录级比较（而非字符串前缀），防止
    output_evil/ 这类前缀兄弟目录绕过白名单；Windows 下 normcase
    统一大小写与斜杠。
    """
    allowed_dirs = [
        os.path.join(SCRIPT_DIR, "output"),
    ]
    abs_path = os.path.normcase(os.path.abspath(output_path))
    for allowed in allowed_dirs:
        allowed = os.path.normcase(os.path.abspath(allowed))
        try:
            if os.path.commonpath([abs_path, allowed]) == allowed:
                return
        except ValueError:
            continue  # 不同盘符，commonpath 无法比较
    raise OutputPathNotAllowedError(f"输出路径不在白名单内: {output_path}")


def _resolve_style(style_name):
    """解析风格名称，返回 styles.json 中对应的配置。"""
    styles_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "styles.json")
    if os.path.exists(styles_path):
        with open(styles_path, "r", encoding="utf-8") as f:
            styles = json.load(f)
        return styles.get(style_name, styles.get("enterprise", {}))
    return {}


class Renderer:
    """统一渲染器核心类。"""

    def __init__(self, spec_path, client_industry=None):
        """初始化渲染器。"""
        if os.environ.get("_PRESALES_CLI_INVOKED") != "1":
            raise NotInvokedViaCLIError("必须通过 python _cli.py 调用，禁止直接 import _renderer")

        self.spec_path = spec_path
        self.client_industry = client_industry
        self.spec = {}
        self.style_name = "enterprise"
        self.style = {}
        self.pages = []

        if os.path.exists(spec_path):
            with open(spec_path, "r", encoding="utf-8") as f:
                self.spec = yaml.safe_load(f) or {}

        # confirmed 门（单点实现在 _cli_guards.require_confirmed，见 §五 0.6）
        from _cli_guards import require_confirmed
        require_confirmed(self.spec)

        # client_name 守卫：spec 中有 client_name 字段但为空时阻断
        if "client_name" in self.spec and not self.spec.get("client_name"):
            raise RenderBlockedError("spec.client_name 存在但为空，无法生成")

        # Sprint Contract：definition_of_done 字段（v2.9.7）
        # 存在时打印验收标准，供 review 对照检查
        self.definition_of_done = self.spec.get("definition_of_done", [])
        if self.definition_of_done:
            print(f"[Sprint Contract] 验收标准 {len(self.definition_of_done)} 条:")
            for i, item in enumerate(self.definition_of_done, 1):
                print(f"  {i}. {item}")

        self.style_name = self.spec.get("style", "enterprise")
        self.style = _resolve_style(self.style_name)
        # v2 主题包名（§5.1 双形态：str = v2 主题；dict/缺省 = None → legacy）
        _spec_theme = self.spec.get("theme")
        self.v2_theme_name = _spec_theme if isinstance(_spec_theme, str) else None
        self.pages = self.spec.get("pages", [])
        self.format_config = self.spec.get("format", {})

        # 渲染报告：跳过/降级/警告收集（§6.3），调用方经 r.report 读取
        self.report = RenderReport()

        # spec 全量元素校验（§6.2）：错误逐条进 report.warnings，
        # 不阻断——非法元素由渲染层降级/跳过，同样进 report。
        # 打印出口统一在 CLI 的 report.summary()（此处不再逐条 print，避免重复）
        try:
            from _renderer.schema import (
                validate_layout_warnings,
                validate_spec, validate_page_warnings, validate_element_warnings)
            for _e in validate_spec(self.spec):
                self.report.warn(f"[spec校验] {_e}")
            # 页面级容量检查（§七 2.5 第一级 源头限量）：warning 级别，
            # 只进 report.warnings 不阻断，阈值见 schema.PAGE_ELEMENTS_WARN
            for _pi, _page in enumerate(self.pages):
                for _w in validate_page_warnings(_page, page_index=_pi):
                    self.report.warn(f"[spec校验] {_w}")
                # v2 P 系版式：必需构件顺序 + 容量上限（§6，warning 级；
                # 必需构件缺失的 error 已由 validate_spec 收集）
                for _w in validate_layout_warnings(_page, page_index=_pi):
                    self.report.warn(f"[spec校验] {_w}")
                # diagram 数量型容量阈值（间距体系 v1 §六 5）：warning 级别，
                # 阈值取排查实测越界数据 -1（schema.SWIMLANE_LANES_WARN 等）
                for _ei, _elem in enumerate(_page.get("elements", []) or []):
                    for _w in validate_element_warnings(_elem, index=_ei):
                        self.report.warn(f"[spec校验] pages[{_pi}].{_w}")
        except Exception as e:  # 校验自身异常也不阻断生成
            self.report.warn(f"spec 校验异常: {e}")

        # outline-to-spec 提取失败占位页：占位 text 是合法元素，渲染 PASS 但
        # 内容缺失——进 report.warnings 提醒，不阻断。
        # 标记串生成点：_outline_to_spec.py 的失败占位 text 元素
        for _pi, _page in enumerate(self.pages, 1):
            for _elem in _page.get("elements", []) or []:
                if (isinstance(_elem, dict)
                        and _EXTRACT_FAIL_MARKER in str(_elem.get("content", ""))):
                    self.report.warn(f"第 {_pi} 页是提取失败占位页，内容缺失")
                    break

    def _has_diagram_elements(self):
        """spec 是否含 diagram 元素（决定是否注入 diagram CSS）。"""
        for page in self.pages:
            for elem in page.get("elements", []):
                if elem.get("type") == "diagram":
                    return True
        return False

    def _has_chrome_elements(self):
        """spec 是否含 v2 页面构件或 flow_rows 图（决定是否注入 v2 CSS）。"""
        from _renderer.page_chrome import CHROME_RENDERERS
        for page in self.pages:
            for elem in page.get("elements", []):
                if elem.get("type") in CHROME_RENDERERS or elem.get("type") == "topnav":
                    return True
                if elem.get("type") == "diagram" and (
                        elem.get("subtype") == "flow_rows"
                        or elem.get("exhibit") or elem.get("source")):
                    return True
        return False

    @staticmethod
    def _hex_to_rgba(hex_color, opacity):
        """#RRGGBB + 不透明度 -> rgba() 字符串。"""
        h = hex_color.lstrip('#')
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return f"rgba({r},{g},{b},{opacity})"

    def _render_cover_html(self, document):
        """渲染 T1 暗色照片封面（spec.document.cover.template == 'dark-photo'）。

        spec 字段：
          background_image  背景照片（相对 HTML 的路径）
          logo_image        logo（深底用反白稿）
          veil / veil_opacity  罩层色与不透明度（缺省 #0A1540 / 0.8）
          confidential      右上机密标注（空则不显示）
          show_date + date  右下日期（date 缺省取 spec.date）
        版式依据视觉基准样板 §5 T1：logo (3%,6%)，标题 x8%/y42%，
        右上 Sensitivity，右下斜体日期。画布固定 1280x720（html-to-ppt 一页）。
        """
        cover = document.get("cover", {}) or {}
        title = document.get("title", "")
        subtitle = document.get("subtitle", "")
        author = document.get("author") or self.spec.get("author", "")
        veil = cover.get("veil", "#0A1540")
        veil_rgba = self._hex_to_rgba(veil, cover.get("veil_opacity", 0.8))
        bg = cover.get("background_image", "")
        logo = cover.get("logo_image", "")
        date_str = cover.get("date") or self.spec.get("date", "")

        parts = ['<div class="cover" style="position:relative;width:1280px;height:720px;'
                 f'overflow:hidden;background:{veil};">']
        if bg:
            parts.append(f'<img src="{_esc(bg)}" style="position:absolute;left:0;top:0;'
                         'width:1280px;height:720px;object-fit:cover;">')
        parts.append(f'<div style="position:absolute;left:0;top:0;width:1280px;height:720px;'
                     f'background:{veil_rgba};"></div>')
        # 内容层：流式布局（absolute 在 dom-to-pptx 下会文本换行重叠）；
        # bg/veil/logo 三个 img 在 PPT 转换时降级丢弃，PPT 得满版 veil 色封面。
        parts.append('<div style="position:relative;height:720px;box-sizing:border-box;'
                     'padding:43px 38px 0 38px;color:#FFFFFF;">')
        parts.append('<div style="display:flex;justify-content:space-between;align-items:flex-start;">')
        if logo:
            parts.append(f'<img src="{_esc(logo)}" style="height:56px;">')
        else:
            parts.append('<div></div>')
        if cover.get("confidential"):
            parts.append(f'<div style="color:#FFFFFF;font-size:12px;opacity:0.9;">{_esc(cover["confidential"])}</div>')
        parts.append('</div>')
        parts.append(f'<div style="margin-top:190px;color:#FFFFFF;font-size:52px;'
                     f'font-weight:700;line-height:1.25;">{_esc(title)}</div>')
        if subtitle:
            parts.append(f'<div style="margin-top:22px;color:rgba(255,255,255,0.92);'
                         f'font-size:22px;line-height:1.5;">{_esc(subtitle)}</div>')
        if cover.get("show_author") and author:
            parts.append(f'<div style="margin-top:14px;color:rgba(255,255,255,0.85);'
                         f'font-size:16px;">{_esc(author)}</div>')
        if cover.get("show_date") and date_str:
            parts.append(f'<div style="margin-top:150px;text-align:right;color:#FFFFFF;'
                         f'font-size:12px;font-style:italic;opacity:0.9;">{_esc(date_str)}</div>')
        parts.append('</div>')
        parts.append('</div>')
        return "\n".join(parts)

    def render_html(self, output_path):
        """生成 HTML 方案文档。"""
        _validate_output_path(output_path)

        document = self.spec.get("document", {})
        title = document.get("title", "方案文档")
        subtitle = document.get("subtitle", "")
        author = document.get("author") or self.spec.get("author", "")
        cover_cfg = document.get("cover", {}) or {}
        cover_html = self._render_cover_html(document) if cover_cfg.get("template") == "dark-photo" else ""

        html_parts = []
        html_parts.append('<!DOCTYPE html>')
        html_parts.append('<html lang="zh-CN">')
        html_parts.append('<head>')
        html_parts.append('    <meta charset="UTF-8">')
        html_parts.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_parts.append(f'    <title>{_esc(title)}</title>')
        html_parts.append('    <style>')
        html_parts.append(self._get_html_styles(has_cover=bool(cover_html)))
        if self._has_diagram_elements():
            from _renderer.diagram import theme as _dg_theme
            html_parts.append(_dg_theme.css_variables(self.style_name))
        if self._has_chrome_elements():
            # v2 页面构件/flow_rows：主题 tokens :root 块 + 构件 CSS（§7/§9）
            from _renderer.diagram import theme as _dg_theme2
            from _renderer.page_chrome import chrome_css
            html_parts.append(_dg_theme2.theme_tokens_css(self.v2_theme_name))
            html_parts.append(chrome_css())
        if self._has_diagram_elements():
            # 旧 diagram 变量桥接：v2 主题下统一色板（桥接块必须后于
            # theme_tokens_css，--t- 变量先定义；legacy 等值零视觉变化）
            from _renderer.diagram import theme as _dg_theme3
            html_parts.append(_dg_theme3.bridge_css(self.v2_theme_name))
        html_parts.append('    </style>')
        html_parts.append('</head>')
        html_parts.append('<body>')
        # v2 规格：P01 hero 已承载封面信息，跳过旧式 h1/subtitle/author 文档头（防重复）
        _has_hero = any(
            isinstance(e, dict) and e.get("type") == "hero"
            for p in self.pages for e in p.get("elements", []))
        if cover_html:
            html_parts.append(cover_html)
        elif not _has_hero:
            html_parts.append(f'    <h1>{_esc(title)}</h1>')
            if subtitle:
                html_parts.append(f'    <p class="subtitle">{_esc(subtitle)}</p>')
            if author:
                html_parts.append(f'    <p class="author">{_esc(author)}</p>')

        from _renderer.schema import PAGE_LAYOUTS
        for page in self.pages:
            page_title = page.get("title", "")
            # v2 P 系版式页：section 骨架 + 页锚点（topnav 跳转），省略 h2
            # 页标题（标题语义由版式必需构件承载：hero / section_tag +
            # action_title，§6）；自由流页（无 layout / 旧值）保持现状零 diff
            is_v2_page = page.get("layout") in PAGE_LAYOUTS
            if is_v2_page:
                pid = _esc(page.get("id", ""))
                html_parts.append(f'    <section class="v2-page" id="{pid}">')
            else:
                html_parts.append(f'    <h2>{_esc(page_title)}</h2>')
            elements = page.get("elements", [])
            for index, elem in enumerate(elements):
                html_parts.append(
                    self._render_element_html(elem, page.get("id"), index))
            if is_v2_page:
                html_parts.append('    </section>')

        html_parts.append('</body>')
        html_parts.append('</html>')

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_parts))

    def _get_html_styles(self, has_cover=False):
        """返回 HTML 样式。

        配色与字体来自样式单一源 _renderer.theme.resolve_theme
        （styles.json 该风格 + spec.theme.colors 逐槽覆盖，缺省保持原配色）：
          theme:
            colors:
              heading / primary / h2 / subtitle / muted /
              table_header_bg / border / card_bg / card_title
        has_cover=True 时放宽 body 宽度以容纳 1280px 封面画布。
        """
        theme_cfg = self.spec.get("theme") or {}
        overrides = theme_cfg.get("colors", {}) if isinstance(theme_cfg, dict) else {}
        theme = resolve_theme(self.style_name, overrides)
        palette = dict(theme.html_colors)
        palette["body_font"] = theme.fonts.get("body", "Microsoft YaHei")
        css = Template("""
            body { font-family: '$body_font', sans-serif; max-width: 1200px; margin: 0 auto; padding: 40px 20px;
                   word-break: normal; overflow-wrap: normal; }
            h1 { color: $heading; border-bottom: 3px solid $primary; padding-bottom: 10px; }
            h2 { color: $h2; margin-top: 30px; }
            .subtitle { color: $subtitle; font-size: 1.2em; }
            .author { color: $muted; font-size: 0.9em; }
            table { border-collapse: collapse; width: 100%; margin: 15px 0; }
            th, td { border: 1px solid $border; padding: 10px; text-align: left; }
            th { background-color: $table_header_bg; }
            ul { margin: 10px 0 18px; padding-left: 20px; }
            li { margin: 5px 0; }
            .bullets { margin: 10px 0 18px; }
            .bullet { margin: 5px 0; padding-left: 18px; text-indent: -18px; }
            .bullet::before { content: "• "; color: $primary; }
            .card { background: $card_bg; border-radius: 8px; padding: 15px; margin: 10px 0; }
            .card-title { font-weight: bold; color: $card_title; }
            .card-body { margin-top: 5px; }
            .pullquote { background: $card_bg; border-left: 4px solid $primary; border-radius: 4px; padding: 14px 18px; margin: 14px 0; }
            .pq-content { font-size: 1.1em; color: $heading; font-weight: 500; line-height: 1.6; }
            .pq-cite { color: $muted; font-size: 0.9em; margin-top: 6px; }
            .phases { display: flex; flex-direction: column; gap: 10px; margin: 12px 0; }
            .phase { background: $card_bg; border-radius: 8px; padding: 12px 16px; }
            .phase-label { font-weight: bold; color: $card_title; }
            .phase-goal { margin-top: 4px; }
        """).substitute(palette)
        if has_cover:
            css += """
            body { max-width: 1320px; padding-top: 0; }
        """
        return css

    def _render_element_html(self, elem, page_id=None, index=None):
        """渲染单个元素为 HTML。

        字段读取走 _renderer.elements normalize 层（§6.1）；不支持的元素
        走 degrade_text 显式降级，未知 type 进 self.report，绝不静默跳过。
        用户文本统一过 _esc 转义（§七 2.6，防注入/防非法 HTML）。
        """
        elem_type, normalized = normalize_element(elem)
        if is_empty_payload(elem_type, normalized):
            # 已知类型但内容为空——多为字段名写错（如 cards 写成 items），
            # 计 skipped 而非静默丢失（通用_schema验收样件的验收语义）
            self.report.skip(page_id, index, elem_type,
                             empty_payload_reason(elem_type))
            return ""
        if elem_type == "text":
            return f'    <p>{_nl2br(normalized["content"])}</p>'
        elif elem_type == "heading":
            # 页面已有 h1（文档标题）/h2（页标题），元素级 heading 从 h3 起
            level = min(normalized["level"] + 2, 6)
            text = normalized["text"]
            return f'    <h{level}>{_esc(text)}</h{level}>'
        elif elem_type == "bullets":
            items = normalized["items"]
            if not items:
                return ""
            # 用 div.bullet 而非 ul/li：dom-to-pptx 对 ul 高度测量有 bug
            # （多行 li 按单行计高，后续元素上移叠印），div 排版定位可靠
            divs = "\n".join(f'        <div class="bullet">{_nl2br(item)}</div>' for item in items)
            return f'    <div class="bullets">\n{divs}\n    </div>'
        elif elem_type == "cards":
            card_html = ""
            for card in normalized["cards"]:
                card_html += '    <div class="card">\n'
                if card["tag"]:
                    card_html += f'        <span class="card-tag">{_esc(card["tag"])}</span>\n'
                card_html += f'        <div class="card-title">{_esc(card["title"])}</div>\n'
                # 字面 \n / 真实换行统一 <br>（_nl2br 内部先 _esc，§七 2.6）
                card_html += f'        <div class="card-body">{_nl2br(card["body"])}</div>\n'
                if card["highlight"]:
                    card_html += f'        <div class="card-highlight">{_esc(card["highlight"])}</div>\n'
                card_html += '    </div>\n'
            return card_html
        elif elem_type == "table":
            # headers/rows 任一缺失已在上方 is_empty_payload 计 skipped
            headers = normalized["headers"]
            rows = normalized["rows"]
            th_html = "\n".join(f'        <th>{_nl2br(h)}</th>' for h in headers)
            rows_html = ""
            for row in rows:
                td_html = "\n".join(f'            <td>{_nl2br(cell)}</td>' for cell in row)
                rows_html += f'        <tr>\n{td_html}\n        </tr>\n'
            return f'    <table>\n        <thead>\n            <tr>\n{th_html}\n            </tr>\n        </thead>\n        <tbody>\n{rows_html}</tbody>\n    </table>'
        elif elem_type == "phases":
            phases_html = '    <div class="phases">\n'
            for phase in normalized["phases"]:
                phases_html += '        <div class="phase">\n'
                phases_html += f'            <div class="phase-label">{_nl2br(phase["name"])}</div>\n'
                phases_html += f'            <div class="phase-goal">{_nl2br(phase["desc"])}</div>\n'
                if phase["actions"]:
                    actions_html = "\n".join(f'                <li>{_esc(a)}</li>' for a in phase["actions"])
                    phases_html += f'            <ul>\n{actions_html}\n            </ul>\n'
                phases_html += '        </div>\n'
            phases_html += '    </div>'
            return phases_html
        elif elem_type == "pullquote":
            pq = f'    <div class="pullquote">\n        <div class="pq-content">{_nl2br(normalized["content"])}</div>\n'
            if normalized["cite"]:
                pq += f'        <div class="pq-cite">{_nl2br(normalized["cite"])}</div>\n'
            pq += '    </div>'
            return pq
        elif elem_type == "architecture_4a":
            # 4A 架构图仅 DOCX 有原生渲染，HTML 端显式降级（§6.1 能力矩阵）
            text = degrade_text(elem_type, normalized, "html")
            self.report.degrade(page_id, index, elem_type, "html", text)
            return f'    <p class="degraded">{text}</p>'
        elif elem_type == "diagram":
            from _renderer.diagram import render_diagram_html
            return render_diagram_html(elem, style=self.style_name)
        elif elem_type == "product_intro_placeholder":
            from _renderer.diagram import render_placeholder_html
            return render_placeholder_html(elem, style=self.style_name)
        # v2 页面构件（§5.2/§7）：HTML 全端 RENDER，主题变量驱动
        from _renderer.page_chrome import render_chrome_html
        # topnav 特殊：合并文档级 brand（logo 槽位，v3.0），元素级声明优先
        if elem_type == "topnav" and isinstance(self.spec.get("brand"), dict):
            merged = dict(normalized)
            merged.setdefault("logo", self.spec["brand"].get("logo", ""))
            merged.setdefault("logo_position",
                              self.spec["brand"].get("logo_position",
                                                     "topnav_left"))
            normalized = merged
        chrome_html = render_chrome_html(elem_type, normalized, pages=self.pages)
        if chrome_html is not None:
            return chrome_html
        # 未知 type 不再静默跳过，进渲染报告（§6.3）
        self.report.skip(page_id, index, elem_type, "HTML 端不支持")
        return ""

    def render_docx(self, output_path):
        """生成 Word 文档（格式感知 + COM 后处理）。"""
        _validate_output_path(output_path)

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 应用格式要求（如果有）
        fmt = self.format_config
        if fmt:
            self._apply_docx_format(doc, fmt)

        document = self.spec.get("document", {})
        title = document.get("title", "方案文档")

        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = document.get("subtitle", "")
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 暗标模式：不写作者/页眉信息
        if not fmt.get("dark_bid", False):
            pass  # 页眉页脚由 COM 后处理完成

        doc.add_paragraph()

        for page in self.pages:
            page_title = page.get("title", "")
            doc.add_heading(page_title, level=1)

            elements = page.get("elements", [])
            for index, elem in enumerate(elements):
                self._render_element_docx(doc, elem, page.get("id"), index)

        doc.save(output_path)

        # COM 后处理（编号/目录/页眉页脚）
        if fmt and (fmt.get("numbering") or fmt.get("toc")):
            try:
                from _renderer.docx_com_postprocess import com_postprocess
                com_postprocess(output_path, fmt)
            except ImportError:
                print("[docx] pywin32 不可用，跳过 COM 后处理（编号/目录需手动）")
            except Exception as e:
                print(f"[docx] COM 后处理失败: {e}")

    def _apply_docx_format(self, doc, fmt):
        """应用招标格式要求到 docx。"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_LINE_SPACING

        body_font = fmt.get("body_font", "仿宋_GB2312")
        body_size = fmt.get("body_size", 16)
        line_spacing = fmt.get("line_spacing", 28)
        first_line_indent = fmt.get("first_line_indent", 2)
        margin = fmt.get("margin", {})

        # 页边距
        for section in doc.sections:
            if margin.get("top"):
                section.top_margin = Cm(margin["top"])
            if margin.get("bottom"):
                section.bottom_margin = Cm(margin["bottom"])
            if margin.get("left"):
                section.left_margin = Cm(margin["left"])
            if margin.get("right"):
                section.right_margin = Cm(margin["right"])

        # 设置 Normal 样式（正文默认样式）
        style = doc.styles["Normal"]
        font = style.font
        font.name = body_font
        font.size = Pt(body_size)
        # 中文字体设置
        from docx.oxml.ns import qn
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), body_font)

        # 行距和首行缩进
        pf = style.paragraph_format
        if line_spacing:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(line_spacing)
        if first_line_indent:
            pf.first_line_indent = Pt(body_size * first_line_indent)

    def _render_element_docx(self, doc, elem, page_id=None, index=None):
        """渲染单个元素到 DOCX。

        字段读取走 _renderer.elements normalize 层（§6.1）；diagram/占位
        元素 DOCX 无原生渲染，输出 degrade_text 降级文本并进 self.report；
        未知 type 进 self.report，不再 else: pass 静默跳过。
        """
        elem_type, normalized = normalize_element(elem)
        if is_empty_payload(elem_type, normalized):
            # 已知类型但内容为空——多为字段名写错，计 skipped 而非静默丢失
            self.report.skip(page_id, index, elem_type,
                             empty_payload_reason(elem_type))
            return
        if elem_type == "text":
            doc.add_paragraph(normalized["content"])
        elif elem_type == "heading":
            doc.add_heading(normalized["text"], level=normalized["level"])
        elif elem_type == "bullets":
            for item in normalized["items"]:
                doc.add_paragraph(item, style="List Bullet")
        elif elem_type == "cards":
            for card in normalized["cards"]:
                p = doc.add_paragraph()
                run = p.add_run(card["title"])
                run.bold = True
                doc.add_paragraph(card["body"])
        elif elem_type == "table":
            # headers/rows 任一缺失已在上方 is_empty_payload 计 skipped
            headers = normalized["headers"]
            rows = normalized["rows"]
            n_cols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=n_cols)
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            for i, row in enumerate(rows):
                if len(row) != n_cols:
                    # schema 校验已报列数不齐；渲染层兜底不崩：
                    # 长出行按 headers 数截断，短出行补空串
                    self.report.warn(
                        f"[渲染] {page_id} elements[{index}] table rows[{i}] "
                        f"列数 {len(row)} 与 headers 列数 {n_cols} 不齐，"
                        f"已按 headers 截断/补齐")
                row_cells = table.rows[i + 1].cells
                for j in range(n_cols):
                    row_cells[j].text = str(row[j]) if j < len(row) else ""
        elif elem_type == "phases":
            for phase in normalized["phases"]:
                p = doc.add_paragraph()
                run = p.add_run(phase["name"])
                run.bold = True
                if phase["desc"]:
                    doc.add_paragraph(phase["desc"])
                for action in phase["actions"]:
                    doc.add_paragraph(action, style="List Bullet")
        elif elem_type == "pullquote":
            # 缩进斜体引文 + 可选署名行（§6.1 补齐 DOCX 端 pullquote）
            from docx.shared import Cm
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            run = p.add_run(normalized["content"])
            run.italic = True
            if normalized["cite"]:
                cite_p = doc.add_paragraph(f'—— {normalized["cite"]}')
                cite_p.paragraph_format.left_indent = Cm(0.75)
        elif elem_type == "architecture_4a":
            layers = elem.get("layers", [])
            for layer in layers:
                name = layer.get("name", "")
                components = layer.get("components", [])
                p = doc.add_paragraph()
                run = p.add_run(name)
                run.bold = True
                for comp in components:
                    doc.add_paragraph(comp, style="List Bullet")
                doc.add_paragraph()  # 层间空行
        elif elem_type in ("diagram", "product_intro_placeholder"):
            # DOCX 无原生渲染，输出显式降级文本（§6.1 能力矩阵）
            text = degrade_text(elem_type, normalized, "docx")
            doc.add_paragraph(text)
            self.report.degrade(page_id, index, elem_type, "docx", text)
        # ---- v2 页面构件 DOCX 端（§5.2 能力矩阵）----
        elif elem_type == "section_tag":
            text = f'{normalized["index"]} · {normalized["label"]}' \
                if normalized["index"] else normalized["label"]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        elif elem_type == "action_title":
            text = "".join(seg["t"] for seg in normalized["segments"])
            doc.add_heading(text, level=2)
            if normalized["sub"]:
                doc.add_paragraph(normalized["sub"])
        elif elem_type == "info_cards":
            # DOCX 走 bullets 段落（§5.2）
            for card in normalized["cards"]:
                p = doc.add_paragraph()
                run = p.add_run(card["title"])
                run.bold = True
                for item in card["items"]:
                    doc.add_paragraph(item, style="List Bullet")
        elif elem_type == "qa_block":
            for item in normalized["items"]:
                p = doc.add_paragraph()
                run = p.add_run(item["q"])
                run.bold = True
                if item["a"]:
                    doc.add_paragraph(item["a"])
        elif elem_type == "evidence_ledger":
            # B-3 证据台账：DOCX 走四列表格（编号/结论/证据/状态）
            if normalized["title"]:
                doc.add_heading(normalized["title"], level=3)
            if normalized["items"]:
                table = doc.add_table(rows=1 + len(normalized["items"]), cols=4)
                hdr = table.rows[0].cells
                for i, h in enumerate(("编号", "结论", "证据", "状态")):
                    hdr[i].text = h
                for ri, it in enumerate(normalized["items"]):
                    cells = table.rows[ri + 1].cells
                    cells[0].text = it["num"] or "—"
                    cells[1].text = it["conclusion"]
                    cells[2].text = it["evidence"]
                    cells[3].text = it["status"] or "—"
        elif elem_type == "risk_register":
            # B-4 风险登记：DOCX 走四列表格（风险/等级/状态/应对）
            if normalized["title"]:
                doc.add_heading(normalized["title"], level=3)
            if normalized["items"]:
                table = doc.add_table(rows=1 + len(normalized["items"]), cols=4)
                hdr = table.rows[0].cells
                for i, h in enumerate(("风险", "等级", "状态", "应对")):
                    hdr[i].text = h
                for ri, it in enumerate(normalized["items"]):
                    cells = table.rows[ri + 1].cells
                    cells[0].text = it["risk"]
                    cells[1].text = it["level"] or "—"
                    cells[2].text = it["status"] or "—"
                    cells[3].text = it["response"]
        elif elem_type == "raci_matrix":
            # B-5 RACI：DOCX 走动态列表格（任务 + 角色列）
            if normalized["title"]:
                doc.add_heading(normalized["title"], level=3)
            roles = normalized["roles"]
            if roles and normalized["tasks"]:
                table = doc.add_table(rows=1 + len(normalized["tasks"]),
                                      cols=1 + len(roles))
                hdr = table.rows[0].cells
                hdr[0].text = "任务"
                for i, r in enumerate(roles):
                    hdr[i + 1].text = r
                for ri, t in enumerate(normalized["tasks"]):
                    cells = table.rows[ri + 1].cells
                    cells[0].text = t["task"]
                    for ci, r in enumerate(roles):
                        cells[ci + 1].text = t["cells"].get(r, "")
        elif elem_type == "decision_board":
            # B-6 决策面板：DOCX 走方案段落 + 推荐
            if normalized["title"]:
                doc.add_heading(normalized["title"], level=3)
            for opt in normalized["options"]:
                p = doc.add_paragraph()
                run = p.add_run(opt["name"])
                run.bold = True
                for pro in opt["pros"]:
                    doc.add_paragraph(f"优：{pro}", style="List Bullet")
                for con in opt["cons"]:
                    doc.add_paragraph(f"劣：{con}", style="List Bullet")
            if normalized["recommendation"]:
                p = doc.add_paragraph()
                run = p.add_run(f"推荐：{normalized['recommendation']}")
                run.bold = True
            if normalized["next_step"]:
                doc.add_paragraph(normalized["next_step"])
        elif elem_type in ("hero", "stat_cards", "kpi_cards", "pain_cards",
                           "legend_bar", "topnav"):
            # DOCX 端 DEGRADE（§5.2），显式降级文本 + report
            text = degrade_text(elem_type, normalized, "docx")
            doc.add_paragraph(text)
            self.report.degrade(page_id, index, elem_type, "docx", text)
        else:
            # 未知 type 不再静默跳过，进渲染报告（§6.3）
            self.report.skip(page_id, index, elem_type, "DOCX 端不支持")


