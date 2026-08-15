# -*- coding: utf-8 -*-
"""DOCX COM 后处理模块（pywin32）。

python-docx 建骨架后：
1. 多级标题编号（一、(一)、1.、1.1、…）：统一用 python-docx 手动写入
   （两条路径共用 _manual_numbering，与是否有 pywin32 无关）
2. 自动生成目录（TOC 域）：有 pywin32 用 Word COM 生成并刷新；
   无 pywin32 插入 TOC 域代码（Word 打开后 Ctrl+A -> F9 更新）
3. 页眉页脚（非暗标时插入项目名/页码）：仅 COM 路径
"""
import os


def com_postprocess(docx_path, fmt_config):
    """COM 后处理主入口。

    Args:
        docx_path: docx 文件路径
        fmt_config: spec.format 字典
    """
    abs_path = os.path.abspath(docx_path)

    try:
        import win32com.client
    except ImportError:
        _fallback_postprocess(docx_path, fmt_config)
        return

    # 多级编号：先用 python-docx 手动写入（原 COM 版 _apply_numbering 是
    # no-op——把 OutlineLevel 读出来再赋回自身，已删除）
    if fmt_config.get("numbering"):
        try:
            from docx import Document
            _doc = Document(abs_path)
            _manual_numbering(_doc, fmt_config.get("numbering_style", "chinese_multi_level"))
            _doc.save(abs_path)
            print("[docx] 多级编号已应用（手动模式）")
        except Exception as e:
            print(f"[docx] 多级编号失败: {e}")

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(abs_path)

        # 自动目录
        if fmt_config.get("toc"):
            _generate_toc(doc, fmt_config.get("toc_levels", "1-3"))

        # 页眉页脚（非暗标）
        if not fmt_config.get("dark_bid", False):
            _add_headers_footers(doc, fmt_config)

        doc.Save()
    finally:
        if doc:
            doc.Close(False)
        if word:
            word.Quit()


def _generate_toc(doc, levels="1-3"):
    """在文档开头插入目录。

    使用 Word COM 的 TablesOfContents.Add 方法。
    """
    try:
        # 跳到文档开头
        rng = doc.Range(0, 0)

        # 插入"目录"标题
        rng.InsertBefore("\n目录\n")
        rng = doc.Range(0, 0)

        # 插入 TOC 域
        # levels 格式 "1-3" -> Word 用 "1-3"
        toc_range = doc.Range(0, 0)
        doc.TablesOfContents.Add(
            Range=toc_range,
            UseHeadingStyles=True,
            UpperHeadingLevel=int(levels.split("-")[0]),
            LowerHeadingLevel=int(levels.split("-")[1]),
        )

        # 更新目录
        for toc in doc.TablesOfContents:
            toc.Update()

        print("[COM] 目录已生成")
    except Exception as e:
        print(f"[COM] 目录生成失败: {e}")


def _add_headers_footers(doc, fmt_config):
    """添加页眉页脚（非暗标模式）。"""
    try:
        for section in doc.Sections:
            # 页眉：项目名
            header = section.Headers(1)  # wdHeaderFooterPrimary = 1
            header.Range.Text = fmt_config.get("project_name", "")

            # 页脚：页码
            footer = section.Footers(1)
            footer.Range.Text = ""
            footer.PageNumbers.Add(
                PageNumberAlignment=2,  # wdAlignPageNumberCenter = 2
                FirstPage=True,
            )

        print("[COM] 页眉页脚已添加")
    except Exception as e:
        print(f"[COM] 页眉页脚失败: {e}")


# ============================================================
# 降级方案（无 pywin32 时）
# ============================================================

def _fallback_postprocess(docx_path, fmt_config):
    """无 pywin32 时，用 python-docx 手动处理。"""
    from docx import Document

    doc = Document(docx_path)

    # 手动编号：给标题加编号文本
    if fmt_config.get("numbering"):
        _manual_numbering(doc, fmt_config.get("numbering_style", "chinese_multi_level"))

    # 手动 TOC：插入 TOC 域代码（Word 打开时按 F9 更新）
    if fmt_config.get("toc"):
        _manual_toc(doc, fmt_config.get("toc_levels", "1-3"))

    doc.save(docx_path)
    print("[docx] 降级后处理完成（编号/目录为手动模式，Word 打开后请按 Ctrl+A -> F9 更新目录）")


def _manual_numbering(doc, style="chinese_multi_level"):
    """手动编号：给标题段落前加编号文本。"""
    counters = [0] * 7  # 7 级编号计数器

    cn_num = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    cn_paren = ["(一)", "(二)", "(三)", "(四)", "(五)", "(六)", "(七)", "(八)", "(九)", "(十)"]

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading ", "").replace("标题", ""))
            except ValueError:
                continue
            if level < 1 or level > 7:
                continue

            # 更新计数器
            counters[level - 1] += 1
            # 重置下级
            for i in range(level, 7):
                counters[i] = 0

            # 生成编号文本
            n = counters[level - 1]
            if style == "chinese_multi_level":
                if level == 1:
                    prefix = f"{cn_num[n - 1] if n <= 10 else n}、"
                elif level == 2:
                    prefix = f"{cn_paren[n - 1] if n <= 10 else f'({n})'} "
                elif level == 3:
                    prefix = f"{n}. "
                elif level == 4:
                    prefix = f"{counters[0]}.{counters[1]}.{counters[2]}.{counters[3]} "
                else:
                    prefix = f"({n}) "
            else:
                prefix = f"{n}. "

            # 在段落开头插入编号
            run = para.add_run(prefix)
            run.bold = True
            # 移到段落最前面：插在第一个 run 之前
            # （原实现 addprevious(para.runs[-1]) 是把自身移到自己前面，no-op，
            #   编号会留在段落末尾）
            if len(para.runs) > 1:
                para.runs[0]._element.addprevious(run._element)


def _manual_toc(doc, levels="1-3"):
    """手动插入 TOC 域代码（Word 打开时需手动更新）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # 在文档开头插入段落
    para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    new_para = para.insert_paragraph_before("目录")
    new_para.style = doc.styles["Heading 1"]

    # 插入 TOC 域
    toc_para = para.insert_paragraph_before("")
    run = toc_para.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    upper, lower = levels.split("-")
    instrText.text = f' TOC \\o "{upper}-{lower}" \\h \\z \\u '

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    fldText = OxmlElement("w:t")
    fldText.text = "[请按 Ctrl+A, F9 更新目录]"

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_separate)
    run._element.append(fldText)
    run._element.append(fldChar_end)
