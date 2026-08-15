# -*- coding: utf-8 -*-
"""DOCX 编号后处理回归测试（P0：编号曾留在段尾 / COM 版为 no-op）。"""

import pytest

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402

from _renderer.docx_com_postprocess import _manual_numbering, com_postprocess  # noqa: E402


def _headings(doc):
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


class TestManualNumbering:
    def _make_doc(self):
        doc = Document()
        doc.add_heading("总体方案", level=1)
        doc.add_heading("项目背景", level=2)
        doc.add_heading("建设目标", level=2)
        doc.add_heading("实施计划", level=1)
        doc.add_heading("阶段安排", level=2)
        return doc

    def test_prefix_at_paragraph_start(self):
        """编号必须在段落开头（回归：曾插在段尾）。"""
        doc = self._make_doc()
        _manual_numbering(doc)
        texts = _headings(doc)
        assert texts[0].startswith("一、")
        assert texts[1].startswith("(一)")
        assert texts[2].startswith("(二)")
        assert texts[3].startswith("二、")
        assert texts[4].startswith("(一)")
        # 原标题文字保留在编号之后
        assert texts[0].endswith("总体方案")

    def test_counters_reset_per_level(self):
        """一级标题递增时二级编号重新计数。"""
        doc = self._make_doc()
        _manual_numbering(doc)
        texts = _headings(doc)
        assert texts[1].startswith("(一)")  # 一、下第一个二级
        assert texts[4].startswith("(一)")  # 二、下重新从 (一) 开始

    def test_com_path_applies_numbering_without_word(self, tmp_path, monkeypatch):
        """com_postprocess 在无 pywin32 时走降级路径且编号真实写入。"""
        import builtins
        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name.startswith("win32com"):
                raise ImportError("no pywin32")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", fake_import)

        doc = self._make_doc()
        path = tmp_path / "t.docx"
        doc.save(str(path))

        com_postprocess(str(path), {"numbering": True})
        doc2 = Document(str(path))
        texts = _headings(doc2)
        assert texts[0].startswith("一、")
        assert texts[3].startswith("二、")
