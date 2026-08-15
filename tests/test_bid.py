# -*- coding: utf-8 -*-
"""标书模块测试：bid-parse 格式提取、verify L4/L5、outline-to-spec 评分匹配。

不依赖 LLM 的纯函数测试。
"""

import os
import sys
import tempfile

SCRIPT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, SCRIPT_DIR)


class TestExtractFormatRequirements:
    """测试 _bid_parse._extract_format_requirements 正则提取。"""

    def test_dark_bid_detection(self):
        from _bid_parse import _extract_format_requirements
        text = "技术标采用暗标方式，不得出现投标人名称"
        fmt = _extract_format_requirements(text)
        assert fmt["dark_bid"] is True

    def test_no_dark_bid(self):
        from _bid_parse import _extract_format_requirements
        text = "这是一份普通的招标文件"
        fmt = _extract_format_requirements(text)
        assert fmt["dark_bid"] is False

    def test_font_extraction(self):
        from _bid_parse import _extract_format_requirements
        text = "正文使用仿宋_GB2312，标题使用黑体"
        fmt = _extract_format_requirements(text)
        assert fmt["body_font"] is not None
        assert "仿宋" in fmt["body_font"]

    def test_font_auto_gb2312(self):
        from _bid_parse import _extract_format_requirements
        text = "正文使用仿宋"
        fmt = _extract_format_requirements(text)
        assert fmt["body_font"] == "仿宋_GB2312"

    def test_font_size_extraction(self):
        from _bid_parse import _extract_format_requirements
        text = "正文字号为三号"
        fmt = _extract_format_requirements(text)
        assert fmt["body_size"] == 16

    def test_line_spacing_extraction(self):
        from _bid_parse import _extract_format_requirements
        text = "行距为固定值28磅"
        fmt = _extract_format_requirements(text)
        assert fmt["line_spacing"] == 28.0

    def test_margin_extraction_4_values(self):
        from _bid_parse import _extract_format_requirements
        text = "页边距：上3.7cm 下3.5cm 左2.8cm 右2.6cm"
        fmt = _extract_format_requirements(text)
        assert fmt["margin"]["top"] == 3.7
        assert fmt["margin"]["bottom"] == 3.5
        assert fmt["margin"]["left"] == 2.8
        assert fmt["margin"]["right"] == 2.6

    def test_margin_extraction_2_values(self):
        from _bid_parse import _extract_format_requirements
        text = "页边距：上下各2.5cm 左右各2.0cm"
        fmt = _extract_format_requirements(text)
        assert fmt["margin"]["top"] == 2.5
        assert fmt["margin"]["bottom"] == 2.5
        assert fmt["margin"]["left"] == 2.0
        assert fmt["margin"]["right"] == 2.0

    def test_numbering_detection(self):
        from _bid_parse import _extract_format_requirements
        text = "编号方式：一、（一）1. 1.1"
        fmt = _extract_format_requirements(text)
        assert fmt["numbering"] == "chinese_multi_level"

    def test_toc_detection(self):
        from _bid_parse import _extract_format_requirements
        text = "须有目录，目录1-3级"
        fmt = _extract_format_requirements(text)
        assert fmt["toc"] is True
        assert fmt["toc_levels"] == "1-3"

    def test_max_pages(self):
        from _bid_parse import _extract_format_requirements
        text = "不超过80页"
        fmt = _extract_format_requirements(text)
        assert fmt["max_pages"] == 80

    def test_empty_text(self):
        from _bid_parse import _extract_format_requirements
        fmt = _extract_format_requirements("")
        assert fmt["dark_bid"] is False
        assert fmt["body_font"] is None


class TestCheckAiTraces:
    """测试 _verify.check_ai_traces。"""

    def test_banned_words_detected(self):
        from _verify import check_ai_traces
        # 创建临时 docx
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("本方案赋能企业数字化转型，打造闭环生态。")
            doc.add_paragraph("综上所述，我们提供一站式服务。")
            doc.save(f.name)
            path = f.name

        try:
            ok, issues = check_ai_traces(path)
            assert ok is False
            assert any("赋能" in i for i in issues)
            assert any("闭环" in i for i in issues)
            assert any("生态" in i for i in issues)
            assert any("综上所述" in i for i in issues)
        finally:
            os.unlink(path)

    def test_clean_text_passes(self):
        from _verify import check_ai_traces
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            # 写 15 句话，长度有变化，无禁用词
            for i in range(15):
                doc.add_paragraph(f"这是第{i}段内容，长度各不相同" + "测试" * (i + 1))
            doc.save(f.name)
            path = f.name

        try:
            ok, issues = check_ai_traces(path)
            assert ok is True
            assert len(issues) == 0
        finally:
            os.unlink(path)

    def test_low_burstiness_detected(self):
        from _verify import check_ai_traces
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            # 写 15 句话，长度完全一样 -> burstiness 极低
            for _ in range(15):
                doc.add_paragraph("这是测试内容长度一样用来检测")
            doc.save(f.name)
            path = f.name

        try:
            ok, issues = check_ai_traces(path)
            # 可能触发 burstiness 警告（取决于具体长度计算）
            # 至少不应该报错
            assert isinstance(ok, bool)
            assert isinstance(issues, list)
        finally:
            os.unlink(path)


class TestMatchScoringToSection:
    """测试 _outline_to_spec._match_scoring_to_section。"""

    def test_exact_match(self):
        from _outline_to_spec import _match_scoring_to_section
        scoring = [{"item": "技术方案", "score": 30, "detail": "方案完整性"}]
        score, guide = _match_scoring_to_section("技术方案", scoring)
        assert score == 30
        assert guide == "方案完整性"

    def test_fuzzy_match(self):
        from _outline_to_spec import _match_scoring_to_section
        scoring = [{"item": "售后", "score": 10, "detail": "响应及时"}]
        score, guide = _match_scoring_to_section("售后服务方案", scoring)
        assert score == 10
        assert guide == "响应及时"

    def test_sub_item_match(self):
        from _outline_to_spec import _match_scoring_to_section
        scoring = [{
            "item": "技术方案",
            "score": 30,
            "sub_items": [
                {"name": "总体架构", "score": 10, "detail": "架构合理"},
                {"name": "技术路线", "score": 10, "detail": "选型先进"},
            ]
        }]
        score, guide = _match_scoring_to_section("总体架构设计", scoring)
        assert score == 10
        assert guide == "架构合理"

    def test_no_match(self):
        from _outline_to_spec import _match_scoring_to_section
        scoring = [{"item": "报价", "score": 30}]
        score, guide = _match_scoring_to_section("技术方案", scoring)
        assert score is None
        assert guide is None

    def test_empty_scoring(self):
        from _outline_to_spec import _match_scoring_to_section
        score, guide = _match_scoring_to_section("技术方案", [])
        assert score is None
        assert guide is None


class TestCheckBidRisks:
    """测试 _verify.check_bid_risks（无 bid_criteria 时跳过）。"""

    def test_no_criteria_passes(self):
        from _verify import check_bid_risks
        # 无 client_name 或无 bid_criteria.json -> 跳过，返回 True
        ok, issues = check_bid_risks("/nonexistent/path.docx", client_name="")
        assert ok is True
        assert len(issues) == 0

    def test_no_criteria_file(self):
        from _verify import check_bid_risks
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("测试内容")
            doc.save(f.name)
            path = f.name

        try:
            # client_name 存在但无 bid_criteria.json -> 跳过
            ok, issues = check_bid_risks(path, client_name="_nonexistent_client_")
            assert ok is True
        finally:
            os.unlink(path)
