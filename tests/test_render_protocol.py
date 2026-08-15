# -*- coding: utf-8 -*-
"""三端渲染器接入元素协议层的测试（重构 Phase 1，§6.1/§6.3）。

覆盖能力矩阵补齐的缺口与 RenderReport 汇聚：
- heading：HTML 出 <h3>-<h6>；PPTD 出大字号粗体主色文本
- pullquote：DOCX 出缩进斜体引文 + "—— {cite}" 署名行
- diagram / product_intro_placeholder：DOCX 出降级文本并进 report.degraded
- architecture_4a：HTML/PPTD 出降级文本并进 report.degraded（DOCX 原生渲染 layers）
- 未知 type：三端都进 report.skipped，不再静默
"""

import os
import shutil
import sys

import pytest
import yaml

SCRIPT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, SCRIPT_DIR)

import _pptd_gen
from _renderer import Renderer, _resolve_style
from _renderer.elements import RenderReport

WORK_DIR = os.path.join(SCRIPT_DIR, "output", "通用", "_test_render_protocol")


@pytest.fixture(autouse=True)
def _cli_env():
    """Renderer 要求 _PRESALES_CLI_INVOKED=1；保存/还原，避免污染其他测试
    （test_diagram.py 等用例会 pop 该变量，不能依赖 import 时设置）。"""
    saved = os.environ.pop("_PRESALES_CLI_INVOKED", None)
    os.environ["_PRESALES_CLI_INVOKED"] = "1"
    try:
        yield
    finally:
        if saved is None:
            os.environ.pop("_PRESALES_CLI_INVOKED", None)
        else:
            os.environ["_PRESALES_CLI_INVOKED"] = saved


def _write_spec(tmp_path, pages):
    """写一份 confirmed spec 到 tmp_path，返回路径。"""
    spec = {
        "confirmed": True,
        "document": {"title": "协议测试文档"},
        "pages": pages,
    }
    spec_path = os.path.join(str(tmp_path), "spec.yml")
    with open(spec_path, "w", encoding="utf-8") as f:
        yaml.safe_dump(spec, f, allow_unicode=True)
    return spec_path


@pytest.fixture
def work_dir():
    """HTML/DOCX 输出目录（_validate_output_path 白名单在 output/ 内）。"""
    os.makedirs(WORK_DIR, exist_ok=True)
    yield WORK_DIR
    shutil.rmtree(WORK_DIR, ignore_errors=True)


def _render_html(spec_path, work_dir):
    r = Renderer(spec_path)
    out = os.path.join(work_dir, "out.html")
    r.render_html(out)
    with open(out, encoding="utf-8") as f:
        return r, f.read()


def _render_docx(spec_path, work_dir):
    r = Renderer(spec_path)
    out = os.path.join(work_dir, "out.docx")
    r.render_docx(out)
    from docx import Document
    return r, Document(out)


def _docx_texts(doc):
    return [p.text for p in doc.paragraphs]


def _build_pptd(spec_path, tmp_path):
    spec = _pptd_gen.load_spec(spec_path)
    style = _resolve_style(spec.get("style", "enterprise"))
    report = RenderReport()
    files, _media = _pptd_gen.build_deck(
        spec, spec_path, style, "deck", os.path.join(str(tmp_path), "deck"),
        report=report)
    return report, files


def _content_text_elements(files):
    """pptd files -> 所有 content 页的 text 元素列表。"""
    result = []
    for name, content in files.items():
        if name.startswith("pages/") and "01_cover" not in name:
            for elem in yaml.safe_load(content)["elements"]:
                if elem.get("elementType") == "text":
                    result.append(elem)
    return result


class TestHeading:
    """heading 补格：HTML 出 h3-h6，PPTD 出大字号粗体主色文本。"""

    def test_html_heading_levels(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "heading", "text": "一级小节", "level": 1},
                {"type": "heading", "text": "四级小节", "level": 4},
                {"type": "heading", "text": "超深小节", "level": 7},
            ],
        }])
        r, html = _render_html(spec_path, work_dir)
        # 页面已有 h1/h2，元素级 heading 从 h3 起，最深钳到 h6
        assert "<h3>一级小节</h3>" in html
        assert "<h6>四级小节</h6>" in html
        assert "<h6>超深小节</h6>" in html
        assert not r.report.has_issues()

    def test_pptd_heading_emits_bold_larger_text(self, tmp_path):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "heading", "text": "小节标题", "level": 2},
            ],
        }])
        report, files = _build_pptd(spec_path, tmp_path)
        headings = [e for e in _content_text_elements(files)
                    if "小节标题" in e["content"]["text"]]
        assert len(headings) == 1
        content = headings[0]["content"]
        assert content["fontSize"] > 14  # 明显大于正文 bodytext 14
        assert content["color"] == "$primary"
        assert "<strong>小节标题</strong>" in content["text"]
        assert headings[0]["elementId"] == "elem-130"  # y 坐标方案同 text
        assert not report.has_issues()


class TestPullquoteDocx:
    """pullquote 补格：DOCX 出缩进斜体引文 + 可选署名行。"""

    def test_quote_with_cite(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "pullquote", "content": "模型是天花板", "cite": "张三"},
            ],
        }])
        r, doc = _render_docx(spec_path, work_dir)
        texts = _docx_texts(doc)
        assert "模型是天花板" in texts
        assert "—— 张三" in texts
        quote_p = next(p for p in doc.paragraphs if p.text == "模型是天花板")
        assert quote_p.runs[0].italic
        assert quote_p.paragraph_format.left_indent is not None
        assert not r.report.has_issues()

    def test_quote_without_cite(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "pullquote", "content": "只有引文"},
            ],
        }])
        r, doc = _render_docx(spec_path, work_dir)
        texts = _docx_texts(doc)
        assert "只有引文" in texts
        assert not any(t.startswith("——") for t in texts)
        assert not r.report.has_issues()


class TestDocxDegrade:
    """diagram / placeholder：DOCX 无原生渲染，出降级文本并进 report。"""

    def test_diagram_degrades_to_text(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p9", "title": "页一",
            "elements": [
                {"type": "text", "content": "前文"},
                {"type": "diagram", "diagram_type": "timeline", "title": "演进图"},
            ],
        }])
        r, doc = _render_docx(spec_path, work_dir)
        assert "[架构图：演进图] 请见 HTML/PPT 版" in _docx_texts(doc)
        assert r.report.degraded == [{
            "page": "p9", "index": 1, "type": "diagram", "target": "docx",
            "message": "[架构图：演进图] 请见 HTML/PPT 版",
        }]

    def test_placeholder_degrades_to_text(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "product_intro_placeholder", "title": "产品介绍"},
            ],
        }])
        r, doc = _render_docx(spec_path, work_dir)
        assert "[产品介绍占位：产品介绍]" in _docx_texts(doc)
        assert len(r.report.degraded) == 1
        assert r.report.degraded[0]["type"] == "product_intro_placeholder"


class TestArchitecture4a:
    """architecture_4a：HTML/PPTD 显式降级，DOCX 原生渲染 layers。"""

    _SPEC_ELEMENTS = [
        {"type": "architecture_4a", "layers": [
            {"name": "业务架构", "components": ["流程梳理"]},
        ]},
    ]

    def test_html_degrades_visibly(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [
            {"id": "p2", "title": "页一", "elements": self._SPEC_ELEMENTS},
        ])
        r, html = _render_html(spec_path, work_dir)
        assert '<p class="degraded">[4A 架构图] 本节内容请见 Word 版</p>' in html
        assert r.report.degraded == [{
            "page": "p2", "index": 0, "type": "architecture_4a",
            "target": "html", "message": "[4A 架构图] 本节内容请见 Word 版",
        }]

    def test_pptd_degrades_to_text(self, tmp_path):
        spec_path = _write_spec(tmp_path, [
            {"id": "p2", "title": "页一", "elements": self._SPEC_ELEMENTS},
        ])
        report, files = _build_pptd(spec_path, tmp_path)
        texts = _content_text_elements(files)
        assert any("[4A 架构图] 本节内容请见 Word 版" in e["content"]["text"]
                   for e in texts)
        assert len(report.degraded) == 1
        assert report.degraded[0]["target"] == "pptd"
        assert report.degraded[0]["page"] == "p2"

    def test_docx_renders_layers_natively(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [
            {"id": "p2", "title": "页一", "elements": self._SPEC_ELEMENTS},
        ])
        r, doc = _render_docx(spec_path, work_dir)
        texts = _docx_texts(doc)
        assert "业务架构" in texts
        assert "流程梳理" in texts
        assert not r.report.degraded


class TestUnknownType:
    """未知 type：三端都进 report.skipped，不再静默跳过。"""

    _PAGES = [{
        "id": "p3", "title": "页一",
        "elements": [
            {"type": "text", "content": "正文"},
            {"type": "tree", "foo": 1},
        ],
    }]

    def test_html_skip_reported(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, self._PAGES)
        r, html = _render_html(spec_path, work_dir)
        assert r.report.skipped == [{
            "page": "p3", "index": 1, "type": "tree", "reason": "HTML 端不支持",
        }]

    def test_docx_skip_reported(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, self._PAGES)
        r, _doc = _render_docx(spec_path, work_dir)
        assert r.report.skipped == [{
            "page": "p3", "index": 1, "type": "tree", "reason": "DOCX 端不支持",
        }]

    def test_pptd_skip_reported(self, tmp_path):
        spec_path = _write_spec(tmp_path, self._PAGES)
        report, _files = _build_pptd(spec_path, tmp_path)
        assert report.skipped == [{
            "page": "p3", "index": 1, "type": "tree", "reason": "PPTD 端不支持",
        }]


class TestReportAggregation:
    """report 汇聚：多种情况混合的 spec，skipped/degraded 计数正确。"""

    _PAGES = [{
        "id": "p1", "title": "混合页",
        "elements": [
            {"type": "text", "content": "正文"},                  # 0 三端正常
            {"type": "architecture_4a", "layers": []},            # 1 html/pptd 降级
            {"type": "diagram", "title": "架构图"},               # 2 docx 降级
            {"type": "tree"},                                     # 3 三端都 skip
        ],
    }]

    def test_html_docx_shared_report(self, tmp_path, work_dir):
        """同一 Renderer 的 report 跨 render_html/render_docx 累计。"""
        spec_path = _write_spec(tmp_path, self._PAGES)
        r = Renderer(spec_path)
        r.render_html(os.path.join(work_dir, "out.html"))
        r.render_docx(os.path.join(work_dir, "out.docx"))
        # tree 在 html/docx 各 skip 一次；arch_4a->html、diagram->docx 各降级一次
        assert len(r.report.skipped) == 2
        assert {s["reason"] for s in r.report.skipped} == \
            {"HTML 端不支持", "DOCX 端不支持"}
        assert len(r.report.degraded) == 2
        targets = sorted((d["type"], d["target"]) for d in r.report.degraded)
        assert targets == [("architecture_4a", "html"), ("diagram", "docx")]
        # §6.2 起 Renderer init 做全量 schema 校验：本 spec 3 个非法元素
        # （arch_4a 空 layers / diagram 缺 diagram_type / tree 未知 type）
        # 各进 1 条 warning，不再静默
        assert "跳过 2 元素 / 降级 2 元素 / 警告 3 条" in r.report.summary()

    def test_pptd_report(self, tmp_path):
        spec_path = _write_spec(tmp_path, self._PAGES)
        report, _files = _build_pptd(spec_path, tmp_path)
        assert len(report.skipped) == 1
        assert report.skipped[0]["type"] == "tree"
        assert len(report.degraded) == 1
        assert report.degraded[0]["type"] == "architecture_4a"

    def test_clean_spec_no_issues(self, tmp_path, work_dir):
        """全部受支持元素的 spec：三端渲染后 report 无任何问题（防误报）。

        元素总高必须装得进内容区（y 130–674）——table 行高 hug 后表头有
        一行实高下限（30.2），2 行表从 72 涨到 95；bullets 只放 1 条
        （类型覆盖不变），否则页底防线对超量页产 warn 是真阳性不是误报。
        """
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "齐全页",
            "elements": [
                {"type": "heading", "text": "小节", "level": 2},
                {"type": "text", "content": "正文"},
                {"type": "bullets", "items": ["甲"]},
                {"type": "cards", "cards": [{"title": "卡", "body": "体"}]},
                {"type": "table", "headers": ["列"], "rows": [["格"]]},
                {"type": "phases", "phases": [
                    {"label": "一期", "goal": "上线", "actions": ["调研"]},
                ]},
                {"type": "pullquote", "content": "引文", "cite": "署名"},
            ],
        }])
        r = Renderer(spec_path)
        r.render_html(os.path.join(work_dir, "out.html"))
        r.render_docx(os.path.join(work_dir, "out.docx"))
        assert not r.report.has_issues()
        report, _files = _build_pptd(spec_path, tmp_path)
        assert not report.has_issues()


class TestTableEmptyPayloadSkipped:
    """table 缺 headers 或 rows：三端统一经 is_empty_payload 进 report.skipped。

    旧行为：HTML 端 `if not headers or not rows: return ""` 静默丢弃；
    pptd 端 rows-only 会渲染无表头表格。
    """

    _CASES = [
        ("rows-only", {"headers": [], "rows": [["a", "b"]]}),
        ("headers-only", {"headers": ["列1", "列2"], "rows": []}),
        ("both-empty", {"headers": [], "rows": []}),
    ]

    @pytest.mark.parametrize("name,table_elem", _CASES, ids=[c[0] for c in _CASES])
    def test_html_docx_skip_reported(self, name, table_elem, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [{"type": "table", **table_elem}],
        }])
        r, html = _render_html(spec_path, work_dir)
        r.render_docx(os.path.join(work_dir, "out.docx"))
        assert "<table>" not in html
        assert len(r.report.skipped) == 2  # HTML/DOCX 各计一次
        for s in r.report.skipped:
            assert s["type"] == "table"
            assert s["reason"] == "表格缺 headers 或 rows"

    @pytest.mark.parametrize("name,table_elem", _CASES, ids=[c[0] for c in _CASES])
    def test_pptd_skip_reported(self, name, table_elem, tmp_path):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [{"type": "table", **table_elem}],
        }])
        report, files = _build_pptd(spec_path, tmp_path)
        assert len(report.skipped) == 1
        assert report.skipped[0]["type"] == "table"
        assert report.skipped[0]["reason"] == "表格缺 headers 或 rows"
        # 不再发出半边 table 元素
        for fname, content in files.items():
            if fname.startswith("pages/"):
                assert "elementType: table" not in content


class TestDocxTableRaggedRow:
    """DOCX 表格行列不齐：渲染层按 headers 截断/补齐兜底不崩，并进 report.warnings。"""

    def test_row_longer_than_headers_truncated(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "table", "headers": ["列1", "列2"],
                 "rows": [["a", "b", "c"], ["x", "y", "z"]]},
            ],
        }])
        r, doc = _render_docx(spec_path, work_dir)  # 不抛 IndexError
        assert len(doc.tables) == 1
        # 按 headers 列数截断：每行第 3 个值丢弃
        for ri in (1, 2):
            assert [c.text for c in doc.tables[0].rows[ri].cells] == \
                [["a", "b"], ["x", "y"]][ri - 1]
        assert any("rows[0]" in w and "截断" in w for w in r.report.warnings)
        assert any("rows[1]" in w and "截断" in w for w in r.report.warnings)

    def test_row_shorter_than_headers_padded(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "table", "headers": ["列1", "列2", "列3"],
                 "rows": [["a"]]},
            ],
        }])
        r, doc = _render_docx(spec_path, work_dir)
        assert [c.text for c in doc.tables[0].rows[1].cells] == ["a", "", ""]
        assert any("rows[0]" in w and "截断" in w for w in r.report.warnings)


class TestExtractFailPlaceholder:
    """outline-to-spec 提取失败占位页：合法元素渲染 PASS 但内容缺失，进 warnings。"""

    _MARKER_CONTENT = "（本章节内容提取失败，请重新运行 outline-to-spec 或手动补充）"

    def test_placeholder_page_warned(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [
            {"id": "p1", "title": "正常页",
             "elements": [{"type": "text", "content": "正常内容"}]},
            {"id": "p2", "title": "失败页",
             "elements": [{"type": "text", "role": "body",
                           "content": self._MARKER_CONTENT}]},
        ])
        r, html = _render_html(spec_path, work_dir)
        # 占位元素照常渲染（不阻断），但 warnings 有记录
        assert self._MARKER_CONTENT in html
        assert any("第 2 页是提取失败占位页" in w for w in r.report.warnings)

    def test_normal_spec_no_false_positive(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [
            {"id": "p1", "title": "页一",
             "elements": [{"type": "text", "content": "正常内容"}]},
        ])
        r, _html = _render_html(spec_path, work_dir)
        assert not any("提取失败占位页" in w for w in r.report.warnings)
        assert not r.report.has_issues()


class TestInitWarningsNotPrinted:
    """§6.2 校验 warnings 统一经 CLI report.summary() 输出，init 不再逐条 print。"""

    def test_no_duplicate_print(self, tmp_path, capsys):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [{"type": "tree"}],  # 未知 type → 1 条 schema warning
        }])
        r = Renderer(spec_path)
        out = capsys.readouterr().out
        assert "[spec校验]" not in out
        # warnings 仍进 report（CLI 摘要统一出口）
        assert any("[spec校验]" in w for w in r.report.warnings)


class TestMultilineHtmlBr:
    """多行文本 <br>：bullets/table cell/phases/pullquote 不再塌陷为空格。"""

    def test_multiline_elements_render_br(self, tmp_path, work_dir):
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "bullets", "items": ["第一行\n第二行"]},
                {"type": "table", "headers": ["列"], "rows": [["甲\n乙"]]},
                {"type": "phases", "phases": [
                    {"name": "一期\n启动", "desc": "上线\n验收"},
                ]},
                {"type": "pullquote", "content": "引文\n折行", "cite": "张三\n署"},
            ],
        }])
        r, html = _render_html(spec_path, work_dir)
        assert "第一行<br>第二行" in html
        assert "<td>甲<br>乙</td>" in html
        assert "一期<br>启动" in html
        assert "上线<br>验收" in html
        assert "引文<br>折行" in html
        assert "张三<br>署" in html

    def test_literal_backslash_n_in_table_cell(self, tmp_path, work_dir):
        r"""table cell 含字面 \n（YAML 单引号写法）同样转 <br>。"""
        spec_path = _write_spec(tmp_path, [{
            "id": "p1", "title": "页一",
            "elements": [
                {"type": "table", "headers": ["列"], "rows": [["甲\\n乙"]]},
            ],
        }])
        r, html = _render_html(spec_path, work_dir)
        assert "<td>甲<br>乙</td>" in html
